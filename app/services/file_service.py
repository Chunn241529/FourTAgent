import PyPDF2
from docx import Document
import pandas as pd
import io
import base64
import os
from typing import Dict, Any, Union, List
import logging
from concurrent.futures import ThreadPoolExecutor

logger = logging.getLogger(__name__)
executor = ThreadPoolExecutor(max_workers=4)


class FileService:

    @staticmethod
    def is_image_file(file) -> bool:
        """Kiểm tra xem file có phải là image không"""
        if hasattr(file, "filename"):
            filename = file.filename.lower()
            return any(
                filename.endswith(ext)
                for ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp"]
            )
        elif isinstance(file, str):
            return file.startswith("data:image")
        return False

    @staticmethod
    def get_file_bytes(file) -> bytes:
        """Lấy bytes từ file object"""
        if hasattr(file, "file"):
            # UploadFile object
            return file.file.read()
        elif isinstance(file, str):
            # Log the start of string to see format
            prefix = file[:50] if len(file) > 50 else file
            logger.info(f"Processing string file input. Prefix: {prefix}")

            if file.startswith("data:"):
                try:
                    header, data = file.split(",", 1)
                    logger.info(f"Decoding base64 file with header: {header}")
                    decoded = base64.b64decode(data)
                    logger.info(f"Decoded {len(decoded)} bytes from base64 string")
                    return decoded
                except Exception as e:
                    logger.error(f"Base64 decoding failed: {e}")
                    # Try direct decode if split fails
                    return base64.b64decode(file)
            else:
                # Maybe it's raw base64 without prefix?
                try:
                    logger.info(
                        "String does not start with data:, attempting direct base64 decode..."
                    )
                    decoded = base64.b64decode(file)
                    logger.info(f"Direct base64 decode success: {len(decoded)} bytes")
                    return decoded
                except:
                    logger.info(
                        "Direct base64 failed, treating as utf-8 string path/content"
                    )
                    return file.encode("utf-8")
        return file

    @staticmethod
    def extract_excel_metadata(file_content: bytes) -> Dict[str, Any]:
        """Trích xuất metadata từ Excel file"""
        try:
            excel_file = pd.ExcelFile(io.BytesIO(file_content))
            metadata = {
                "file_type": "excel",
                "sheet_count": len(excel_file.sheet_names),
                "sheets": [],
            }

            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    sheet_info = {
                        "name": sheet_name,
                        "rows": len(df),
                        "columns": len(df.columns),
                        "columns_list": df.columns.tolist(),
                        "sample_data": (
                            df.head(3).to_dict("records") if not df.empty else []
                        ),
                    }
                    metadata["sheets"].append(sheet_info)
                except Exception as e:
                    logger.warning(f"Failed to analyze sheet {sheet_name}: {e}")
                    continue

            return metadata
        except Exception as e:
            logger.error(f"Error extracting Excel metadata: {e}")
            return {}

    @staticmethod
    def extract_docx_metadata(file_content: bytes) -> Dict[str, Any]:
        """Trích xuất metadata từ DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            metadata = {
                "file_type": "docx",
                "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                "table_count": len(doc.tables),
                "sections_count": len(doc.sections),
            }
            return metadata
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            return {}

    @staticmethod
    def extract_text_from_file(file_content: Union[bytes, str]) -> str:
        """Trích xuất text từ nhiều định dạng file"""
        if isinstance(file_content, str):
            try:
                file_content = base64.b64decode(file_content)
            except:
                return file_content

        def sync_extract(file_content: bytes) -> str:
            # PDF extraction
            text = FileService._extract_pdf_text(file_content)
            if text:
                return text

            # CSV extraction
            text = FileService._extract_csv_text(file_content)
            if text:
                return text

            # Excel extraction
            text = FileService._extract_excel_text(file_content)
            if text:
                return text

            # DOCX extraction
            text = FileService._extract_docx_text(file_content)
            if text:
                return text

            # TXT extraction
            text = FileService._extract_txt_text(file_content)
            if text:
                return text

            # Fallback decoding
            return FileService._fallback_decode(file_content)

        return executor.submit(sync_extract, file_content).result()

    @staticmethod
    def _extract_pdf_text(file_content: bytes) -> str:
        """Extract text from PDF"""
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = "\n".join(
                page.extract_text() for page in reader.pages if page.extract_text()
            )
            if text.strip():
                logger.info("Successfully extracted text from PDF")
                return text[:20000]
        except Exception as e:
            logger.warning(f"PDF extraction failed: {e}")
        return ""

    @staticmethod
    def _extract_csv_text(file_content: bytes) -> str:
        """Extract text from CSV"""
        try:
            df = pd.read_csv(io.BytesIO(file_content))
            text = df.to_string(index=False)
            if text.strip():
                logger.info("Successfully extracted text from CSV")
                return text[:20000]
        except Exception as e:
            logger.warning(f"CSV extraction failed: {e}")
        return ""

    @staticmethod
    def _extract_excel_text(file_content: bytes) -> str:
        """Extract text từ Excel (multi-sheet)"""
        try:
            excel_file = pd.ExcelFile(io.BytesIO(file_content))
            all_sheets_text = []

            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    df = df.fillna("")

                    sheet_header = f"=== SHEET: {sheet_name} ===\n"
                    sheet_header += f"Columns: {', '.join(df.columns.astype(str))}\n"
                    sheet_header += (
                        f"Shape: {len(df)} rows x {len(df.columns)} columns\n"
                    )
                    sheet_header += "-" * 50 + "\n"

                    sheet_text = df.to_string(index=False, max_rows=100)
                    all_sheets_text.append(sheet_header + sheet_text)

                    logger.info(
                        f"Extracted sheet: {sheet_name} with {len(df)} rows, {len(df.columns)} columns"
                    )

                except Exception as e:
                    logger.warning(f"Failed to extract sheet {sheet_name}: {e}")
                    all_sheets_text.append(
                        f"=== SHEET: {sheet_name} ===\n[Error extracting this sheet: {e}]"
                    )
                    continue

            if all_sheets_text:
                combined_text = "\n\n".join(all_sheets_text)
                logger.info(
                    f"Successfully extracted {len(all_sheets_text)} sheets from Excel"
                )
                return combined_text[:20000]

        except Exception as e:
            logger.warning(f"Excel extraction failed: {e}")
        return ""

    @staticmethod
    def _extract_docx_text(file_content: bytes) -> str:
        """Extract text từ DOCX"""
        try:
            doc = Document(io.BytesIO(file_content))
            all_text = []

            # Extract paragraphs
            paragraph_count = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text)
                    paragraph_count += 1

            # Extract tables
            table_count = 0
            for i, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                    if row_data:
                        table_data.append(" | ".join(row_data))

                if table_data:
                    table_header = f"\n--- TABLE {i} ---"
                    all_text.append(table_header)
                    all_text.extend(table_data)
                    table_count += 1

            if all_text:
                combined_text = "\n".join(all_text)
                logger.info(
                    f"Successfully extracted {paragraph_count} paragraphs and {table_count} tables from DOCX"
                )
                return combined_text[:20000]

        except Exception as e:
            logger.warning(f"DOCX extraction failed: {e}")
        return ""

    @staticmethod
    def _extract_txt_text(file_content: bytes) -> str:
        """Extract text từ TXT"""
        try:
            text = file_content.decode("utf-8", errors="replace")
            if text.strip():
                logger.info("Successfully extracted text from TXT")
                return text[:20000]
        except Exception as e:
            logger.warning(f"TXT extraction failed: {e}")
        return ""

    @staticmethod
    def _fallback_decode(file_content: bytes) -> str:
        """Fallback decoding với multiple encodings"""
        for encoding in ["utf-8", "latin-1", "cp1252", "iso-8859-1", "utf-16"]:
            try:
                text = file_content.decode(encoding, errors="replace")
                if len(text.strip()) > 100:
                    logger.info(f"Successfully decoded text with {encoding}")
                    return text[:20000]
            except:
                continue

        logger.warning("Could not extract meaningful text from file")
        return ""

    @staticmethod
    def process_file_for_chat(file, user_id: int, conversation_id: int) -> str:
        """Xử lý file cho chat (trả về context string)"""
        if not file:
            return ""

        file_bytes = FileService.get_file_bytes(file)
        filename = getattr(file, "filename", "uploaded_file")

        if FileService.is_image_file(file):
            return ""  # Image được xử lý riêng

        # Xử lý RAG cho non-image files
        from app.services.rag_service import RAGService

        return RAGService.process_file_for_rag(
            file_bytes, user_id, conversation_id, filename
        )
